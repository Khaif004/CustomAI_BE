from __future__ import annotations

import io
import json
import logging
import re
from typing import List, Optional

import aiohttp
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.auth.security import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/documents", tags=["documents"])

class DocGenRequest(BaseModel):
    topic: str = Field(..., min_length=3, max_length=2000)
    doc_type: str = Field(..., pattern=r"^(word|pdf|excel)$")
    additional_context: Optional[str] = Field(None, max_length=4000)
    app_id: Optional[str] = Field(None, pattern=r"^[a-zA-Z0-9_-]*$")

async def _get_document_content(topic: str, doc_type: str, additional_context: Optional[str]) -> dict:
    """Call the SAP AI Core / OpenAI LLM to produce structured document content as JSON."""
    from app.config import get_settings
    settings = get_settings()

    context_block = f"\n\nAdditional context:\n{additional_context}" if additional_context else ""

    if doc_type == "excel":
        schema_hint = """Return a JSON object with this structure:
{
  "title": "Document title",
  "sheets": [
    {
      "name": "Sheet name (max 31 chars)",
      "description": "What this sheet contains",
      "headers": ["Col1", "Col2", ...],
      "rows": [
        ["value1", "value2", ...],
        ...
      ],
      "summary": "Optional brief summary row (null if not needed)"
    }
  ]
}
Include at least 2 sheets when it makes sense. Populate each sheet with realistic, detailed data rows (at least 8-15 rows per sheet). Use clear, professional column names."""
    else:
        schema_hint = """Return a JSON object with this structure:
{
  "title": "Document title",
  "subtitle": "Optional subtitle or null",
  "sections": [
    {
      "heading": "Section heading",
      "level": 1,
      "content": "Paragraph text. Can be multiple sentences.",
      "bullets": ["bullet point 1", "bullet point 2"],
      "table": {
        "headers": ["Col1", "Col2", "Col3"],
        "rows": [["v1","v2","v3"], ...]
      }
    }
  ],
  "conclusion": "Concluding paragraph or null"
}
Rules:
- Include 5-8 well-structured sections with headings.
- Use 'bullets' for lists (can be null if not applicable).
- Use 'table' only when tabular data is genuinely useful (can be null).
- Content should be detailed, professional and comprehensive (3-5 sentences per section minimum).
- level 1 = main section heading, level 2 = sub-section."""

    prompt = f"""You are a professional technical writer. Generate a comprehensive, high-quality document about the following topic for a {doc_type.upper()} file.

Topic: {topic}{context_block}

{schema_hint}

IMPORTANT: Return ONLY valid JSON. No markdown fences, no explanatory text outside the JSON."""

    if settings.llm_provider == "sap_ai_core":
        return await _call_sap_ai_core(prompt, settings)
    else:
        return await _call_openai(prompt, settings)


async def _call_sap_ai_core(prompt: str, settings) -> dict:
    from app.agents.sap_ai_core_agent import SAPAICoreAuth
    auth = SAPAICoreAuth(
        auth_url=settings.sap_aicore_auth_url,
        client_id=settings.sap_aicore_client_id,
        client_secret=settings.sap_aicore_client_secret,
    )
    token = await auth.get_token()
    inference_url = f"{settings.sap_aicore_url.rstrip('/')}/v2/inference/deployments/{settings.sap_aicore_deployment_id}/completion"

    payload = {
        "orchestration_config": {
            "module_configurations": {
                "llm_module_config": {
                    "model_name": settings.sap_aicore_model_id,
                    "model_params": {"max_tokens": 4096, "temperature": 0.3},
                },
                "templating_module_config": {
                    "template": [
                        {"role": "system", "content": "You are a professional technical writer. Always return only valid JSON."},
                        {"role": "user", "content": "{{?prompt}}"},
                    ]
                },
            }
        },
        "input_params": {"prompt": prompt},
    }

    async with aiohttp.ClientSession() as session:
        async with session.post(
            inference_url,
            json=payload,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json", "AI-Resource-Group": settings.sap_aicore_resource_group or "default"},
            timeout=aiohttp.ClientTimeout(total=120),
        ) as resp:
            text = await resp.text()
            if resp.status != 200:
                raise Exception(f"AI Core error {resp.status}: {text[:400]}")
            result = json.loads(text)
            content = result["module_results"]["llm"]["choices"][0]["message"]["content"]
            return _parse_json_response(content)


async def _call_openai(prompt: str, settings) -> dict:
    from langchain_openai import ChatOpenAI
    from langchain_core.messages import HumanMessage, SystemMessage
    llm = ChatOpenAI(model=getattr(settings, "openai_model", "gpt-4"), temperature=0.3, api_key=settings.openai_api_key, max_tokens=4096)
    msgs = [SystemMessage(content="You are a professional technical writer. Always return only valid JSON."), HumanMessage(content=prompt)]
    response = await llm.ainvoke(msgs)
    return _parse_json_response(response.content)


def _parse_json_response(text: str) -> dict:
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return json.loads(text)

def _build_word(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(data.get("title", "Document"))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    if data.get("subtitle"):
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(data["subtitle"])
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        sub_run.italic = True

    doc.add_paragraph()  # spacer

    def _add_table(doc, headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1A56DB")
            tcPr.append(shd)
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val) if val is not None else ""
        doc.add_paragraph()

    for section in data.get("sections", []):
        level = section.get("level", 1)
        heading_style = "Heading 1" if level == 1 else "Heading 2"
        doc.add_heading(section.get("heading", ""), level=level)

        if section.get("content"):
            para = doc.add_paragraph(section["content"])
            para.paragraph_format.space_after = Pt(6)

        for bullet in section.get("bullets") or []:
            doc.add_paragraph(bullet, style="List Bullet")

        tbl = section.get("table")
        if tbl and tbl.get("headers") and tbl.get("rows"):
            _add_table(doc, tbl["headers"], tbl["rows"])

    if data.get("conclusion"):
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(data["conclusion"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf(data: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )

    BLUE = colors.HexColor("#1A56DB")
    LIGHT_BLUE = colors.HexColor("#EFF6FF")
    DARK = colors.HexColor("#111827")
    GRAY = colors.HexColor("#6B7280")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("DocTitle", parent=styles["Normal"], fontSize=22, textColor=BLUE, spaceAfter=6, alignment=TA_CENTER, fontName="Helvetica-Bold")
    subtitle_style = ParagraphStyle("DocSub", parent=styles["Normal"], fontSize=12, textColor=GRAY, spaceAfter=16, alignment=TA_CENTER, fontName="Helvetica-Oblique")
    h1_style = ParagraphStyle("H1", parent=styles["Normal"], fontSize=14, textColor=BLUE, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold")
    h2_style = ParagraphStyle("H2", parent=styles["Normal"], fontSize=12, textColor=DARK, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, textColor=DARK, spaceAfter=8, leading=15, alignment=TA_JUSTIFY, fontName="Helvetica")
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, textColor=DARK, spaceAfter=4, leftIndent=16, bulletIndent=4, leading=14, fontName="Helvetica")

    story = []

    story.append(Paragraph(data.get("title", "Document"), title_style))
    if data.get("subtitle"):
        story.append(Paragraph(data["subtitle"], subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLUE, spaceAfter=14))

    def add_table(headers, rows):
        col_count = len(headers)
        avail_width = A4[0] - 5.5*cm
        col_w = avail_width / col_count
        tbl_data = [[Paragraph(f"<b>{h}</b>", ParagraphStyle("th", fontSize=9, textColor=colors.white, fontName="Helvetica-Bold")) for h in headers]]
        for row in rows:
            tbl_data.append([Paragraph(str(v) if v is not None else "", ParagraphStyle("td", fontSize=9, fontName="Helvetica")) for v in row])
        tbl = Table(tbl_data, colWidths=[col_w]*col_count, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_BLUE]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 10))

    for section in data.get("sections", []):
        level = section.get("level", 1)
        h_style = h1_style if level == 1 else h2_style
        block = [Paragraph(section.get("heading", ""), h_style)]

        if section.get("content"):
            block.append(Paragraph(section["content"], body_style))

        for bullet in section.get("bullets") or []:
            block.append(Paragraph(f"• {bullet}", bullet_style))

        tbl = section.get("table")
        if tbl and tbl.get("headers") and tbl.get("rows"):
            story.extend(block)
            block = []
            add_table(tbl["headers"], tbl["rows"])

        if block:
            story.append(KeepTogether(block))

    if data.get("conclusion"):
        story.append(Paragraph("Conclusion", h1_style))
        story.append(Paragraph(data["conclusion"], body_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _build_excel(data: dict) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, GradientFill
    from openpyxl.utils import get_column_letter

    BLUE = "1A56DB"
    LIGHT_BLUE = "EFF6FF"
    HEADER_FONT_COLOR = "FFFFFF"
    ALT_ROW = "F3F4F6"

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for sheet_data in data.get("sheets", []):
        name = (sheet_data.get("name") or "Sheet")[:31]
        ws = wb.create_sheet(title=name)

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(sheet_data.get("headers", [])), 1))
        title_cell = ws.cell(row=1, column=1, value=data.get("title", "") + " — " + name)
        title_cell.font = Font(bold=True, size=13, color=HEADER_FONT_COLOR)
        title_cell.fill = PatternFill("solid", fgColor=BLUE)
        title_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[1].height = 28

        if sheet_data.get("description"):
            desc_cell = ws.cell(row=2, column=1, value=sheet_data["description"])
            desc_cell.font = Font(italic=True, size=10, color="6B7280")
            desc_cell.alignment = Alignment(horizontal="left")
            ws.row_dimensions[2].height = 18
            data_start_row = 4
        else:
            data_start_row = 3

        headers = sheet_data.get("headers", [])
        # Header row
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=data_start_row, column=col_idx, value=header)
            cell.font = Font(bold=True, color=HEADER_FONT_COLOR, size=10)
            cell.fill = PatternFill("solid", fgColor="374151")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        ws.row_dimensions[data_start_row].height = 22

        for row_idx, row in enumerate(sheet_data.get("rows", []), start=data_start_row + 1):
            fill_color = ALT_ROW if (row_idx - data_start_row) % 2 == 0 else "FFFFFF"
            for col_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.fill = PatternFill("solid", fgColor=fill_color)
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                cell.border = border
                cell.font = Font(size=10)
            ws.row_dimensions[row_idx].height = 18

        if sheet_data.get("summary"):
            sum_row = data_start_row + len(sheet_data.get("rows", [])) + 2
            summary = sheet_data["summary"]
            if isinstance(summary, list):
                for col_idx, val in enumerate(summary, start=1):
                    sc = ws.cell(row=sum_row, column=col_idx, value=val if not isinstance(val, (dict, list)) else str(val))
                    sc.font = Font(bold=True, italic=True, size=10, color="374151")
            else:
                sum_cell = ws.cell(row=sum_row, column=1, value=str(summary) if summary is not None else "")
                sum_cell.font = Font(bold=True, italic=True, size=10, color="374151")

        for col_idx in range(1, len(headers) + 1):
            max_len = max(
                (len(str(ws.cell(row=r, column=col_idx).value or "")) for r in range(data_start_row, data_start_row + len(sheet_data.get("rows", [])) + 2)),
                default=10,
            )
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 3, len(str(headers[col_idx - 1])) + 4), 45)

        ws.freeze_panes = ws.cell(row=data_start_row + 1, column=1)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

MIME_TYPES = {
    "word":  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":   "application/pdf",
    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
EXTENSIONS = {"word": "docx", "pdf": "pdf", "excel": "xlsx"}
BUILDERS = {"word": _build_word, "pdf": _build_pdf, "excel": _build_excel}


@router.post("/generate")
async def generate_document(request: DocGenRequest, current_user=Depends(get_current_user)):
    """
    Generate a formatted document (Word / PDF / Excel) on any topic using the LLM.
    Returns the file as a binary download.
    """
    try:
        logger.info(f"Generating {request.doc_type} document: '{request.topic[:80]}'")
        data = await _get_document_content(request.topic, request.doc_type, request.additional_context)
        file_bytes = BUILDERS[request.doc_type](data)

        safe_name = re.sub(r"[^\w\-]", "_", (data.get("title") or request.topic)[:50])
        filename = f"{safe_name}.{EXTENSIONS[request.doc_type]}"

        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=MIME_TYPES[request.doc_type],
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"Document generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))
